"""
AI Services for Saathi - RAG, transcription, and document processing.
"""

import os
import logging
import tempfile
from typing import Dict, Any, List, Optional
from django.conf import settings
import requests

logger = logging.getLogger(__name__)

# Global services
rag_service = None
transcription_service = None


class RAGService:
    """Retrieval-Augmented Generation service using LlamaIndex and FAISS/Chroma."""
    
    def __init__(self):
        self.index = None
        self.vector_store = None
        self.initialized = False
        
        try:
            self._initialize_rag()
            self.initialized = True
            logger.info("RAG service initialized successfully")
        except Exception as e:
            logger.error(f"Failed to initialize RAG service: {e}")
            self.initialized = False
    
    def _initialize_rag(self):
        """Initialize LlamaIndex with FAISS/Chroma vector store."""
        try:
            from llama_index.core import VectorStoreIndex, Document, StorageContext
            from llama_index.vector_stores.faiss import FaissVectorStore
            from llama_index.embeddings.huggingface import HuggingFaceEmbedding
            import faiss
            
            # Initialize embedding model
            embed_model = HuggingFaceEmbedding(
                model_name="sentence-transformers/all-MiniLM-L6-v2"
            )
            
            # Check if index exists
            index_path = settings.FAISS_INDEX_PATH / "index.faiss"
            
            if index_path.exists():
                # Load existing index
                faiss_index = faiss.read_index(str(index_path))
                vector_store = FaissVectorStore(faiss_index=faiss_index)
                storage_context = StorageContext.from_defaults(vector_store=vector_store)
                self.index = VectorStoreIndex.from_vector_store(
                    vector_store=vector_store,
                    embed_model=embed_model
                )
            else:
                # Create new index
                faiss_index = faiss.IndexFlatL2(384)  # dimension for all-MiniLM-L6-v2
                vector_store = FaissVectorStore(faiss_index=faiss_index)
                storage_context = StorageContext.from_defaults(vector_store=vector_store)
                self.index = VectorStoreIndex(
                    [],
                    storage_context=storage_context,
                    embed_model=embed_model
                )
            
            self.vector_store = vector_store
            
        except ImportError as e:
            logger.error(f"Missing dependencies for RAG: {e}")
            raise
    
    def ingest_document(
        self, 
        file_url: str, 
        uid: str, 
        filename: str = None
    ) -> Dict[str, Any]:
        """Download and ingest a document into the RAG system."""
        
        if not self.initialized:
            return {
                'success': False,
                'error': 'RAG service not initialized',
                'chunks_added': 0
            }
        
        try:
            # Download file from URL
            response = requests.get(file_url, timeout=30)
            response.raise_for_status()
            
            # Determine file type and extract text
            content_type = response.headers.get('content-type', '').lower()
            
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                temp_file.write(response.content)
                temp_path = temp_file.name
            
            try:
                extracted_text = self._extract_text(temp_path, content_type)
                
                if not extracted_text:
                    return {
                        'success': False,
                        'error': 'Could not extract text from document',
                        'chunks_added': 0
                    }
                
                # Create document chunks
                chunks = self._create_chunks(extracted_text, uid, filename or 'uploaded_document')
                
                # Add to index
                from llama_index.core import Document
                documents = [
                    Document(
                        text=chunk['text'],
                        metadata={
                            'uid': uid,
                            'filename': filename,
                            'chunk_id': chunk['id'],
                            'source': 'user_upload'
                        }
                    )
                    for chunk in chunks
                ]
                
                # Insert documents into index
                for doc in documents:
                    self.index.insert(doc)
                
                # Save index
                self._save_index()
                
                return {
                    'success': True,
                    'chunks_added': len(chunks),
                    'extracted_text': extracted_text[:500] + '...' if len(extracted_text) > 500 else extracted_text
                }
                
            finally:
                # Clean up temp file
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
                    
        except Exception as e:
            logger.error(f"Document ingestion error: {e}")
            return {
                'success': False,
                'error': str(e),
                'chunks_added': 0
            }
    
    def _extract_text(self, file_path: str, content_type: str) -> str:
        """Extract text from various file formats."""
        
        try:
            if 'pdf' in content_type:
                return self._extract_pdf_text(file_path)
            elif 'word' in content_type or 'document' in content_type:
                return self._extract_docx_text(file_path)
            elif 'text' in content_type:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                # Try to read as text
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
        except Exception as e:
            logger.error(f"Text extraction error: {e}")
            return ""
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF files."""
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(file_path)
            text = ""
            
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return ""
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX files."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return ""
    
    def _create_chunks(self, text: str, uid: str, filename: str) -> List[Dict[str, Any]]:
        """Create text chunks for vector storage."""
        
        # Simple chunking strategy - split by sentences/paragraphs
        chunks = []
        chunk_size = 512
        overlap = 50
        
        words = text.split()
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk_words = words[i:i + chunk_size]
            chunk_text = ' '.join(chunk_words)
            
            if len(chunk_text.strip()) > 50:  # Minimum chunk size
                chunks.append({
                    'id': f"{uid}_{filename}_{i}",
                    'text': chunk_text,
                    'uid': uid,
                    'filename': filename,
                    'chunk_index': len(chunks)
                })
        
        return chunks
    
    def query(self, query_text: str, uid: str = None, top_k: int = 3) -> List[Dict[str, Any]]:
        """Query the RAG system for relevant documents."""
        
        if not self.initialized:
            return []
        
        try:
            query_engine = self.index.as_query_engine(similarity_top_k=top_k)
            response = query_engine.query(query_text)
            
            results = []
            for node in response.source_nodes:
                # Filter by user if specified
                if uid and node.metadata.get('uid') != uid:
                    continue
                    
                results.append({
                    'text': node.text,
                    'score': node.score,
                    'metadata': node.metadata
                })
            
            return results
            
        except Exception as e:
            logger.error(f"RAG query error: {e}")
            return []
    
    def _save_index(self):
        """Save the FAISS index to disk."""
        try:
            if self.vector_store and self.vector_store.faiss_index:
                import faiss
                index_path = settings.FAISS_INDEX_PATH / "index.faiss"
                faiss.write_index(self.vector_store.faiss_index, str(index_path))
        except Exception as e:
            logger.error(f"Failed to save index: {e}")


class TranscriptionService:
    """Audio transcription service using Whisper API or local models."""
    
    def __init__(self):
        self.whisper_available = bool(settings.OPENAI_API_KEY)
        self.local_whisper = None
        
        if not self.whisper_available:
            try:
                import whisper
                self.local_whisper = whisper.load_model("base")
                logger.info("Local Whisper model loaded")
            except Exception as e:
                logger.warning(f"Could not load local Whisper: {e}")
    
    def transcribe_audio(self, audio_file_path: str) -> Dict[str, Any]:
        """Transcribe audio file to text."""
        
        try:
            if self.whisper_available:
                return self._transcribe_with_openai(audio_file_path)
            elif self.local_whisper:
                return self._transcribe_with_local_whisper(audio_file_path)
            else:
                return {
                    'success': False,
                    'text': '',
                    'error': 'No transcription service available'
                }
        except Exception as e:
            logger.error(f"Transcription error: {e}")
            return {
                'success': False,
                'text': '',
                'error': str(e)
            }
    
    def _transcribe_with_openai(self, audio_file_path: str) -> Dict[str, Any]:
        """Transcribe using OpenAI Whisper API."""
        try:
            import openai
            
            client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)
            
            with open(audio_file_path, 'rb') as audio_file:
                transcript = client.audio.transcriptions.create(
                    model="whisper-1",
                    file=audio_file,
                    response_format="text"
                )
            
            return {
                'success': True,
                'text': transcript.strip(),
                'service': 'openai_whisper'
            }
            
        except Exception as e:
            logger.error(f"OpenAI Whisper error: {e}")
            raise
    
    def _transcribe_with_local_whisper(self, audio_file_path: str) -> Dict[str, Any]:
        """Transcribe using local Whisper model."""
        try:
            result = self.local_whisper.transcribe(audio_file_path)
            
            return {
                'success': True,
                'text': result['text'].strip(),
                'service': 'local_whisper'
            }
            
        except Exception as e:
            logger.error(f"Local Whisper error: {e}")
            raise


def get_rag_service() -> RAGService:
    """Get the global RAG service instance."""
    global rag_service
    if rag_service is None:
        rag_service = RAGService()
    return rag_service


def get_transcription_service() -> TranscriptionService:
    """Get the global transcription service instance."""
    global transcription_service
    if transcription_service is None:
        transcription_service = TranscriptionService()
    return transcription_service


def initialize_ai_services():
    """Initialize all AI services."""
    global rag_service, transcription_service
    
    try:
        rag_service = RAGService()
        transcription_service = TranscriptionService()
        
        logger.info("AI services initialized successfully")
        return True
    except Exception as e:
        logger.error(f"Failed to initialize AI services: {e}")
        return False